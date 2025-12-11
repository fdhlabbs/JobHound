# resume_service.py
from typing import Dict, Any
from pathlib import Path

from docx import Document
from docx.shared import Pt

from ai_service import generate_resume_structure as _ai_generate_resume_structure


def generate_resume_structure(context: Dict[str, Any]) -> Dict[str, Any]:
    """
    Thin wrapper so the rest of the app can call
    resume_service.generate_resume_structure(...)
    but the actual model call lives in ai_service.
    """
    return _ai_generate_resume_structure(context)


def _set_paragraph_font(paragraph, size: int = 11, bold: bool = False):
    # Currently unused, but kept in case you want to style later.
    for run in paragraph.runs:
        font = run.font
        font.size = Pt(size)
        font.bold = bold


def build_resume_docx(resume: Dict[str, Any], out_path: Path | str) -> None:
    """
    Build a simple resume from the structured JSON:

    {
      "meta": {...},
      "header": {...},
      "summary": "...",
      "skills": [...],
      "experience": [...],
      "education": [...],
      "projects": [...],
      "extras": [...]
    }
    """
    # Normalise path & ensure folder exists
    if isinstance(out_path, str):
        out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    meta = resume.get("meta", {}) or {}
    header = resume.get("header", {}) or {}
    summary = resume.get("summary", "") or ""
    skills = resume.get("skills", []) or []
    experience = resume.get("experience", []) or []
    education = resume.get("education", []) or []
    projects = resume.get("projects", []) or []
    extras = resume.get("extras", []) or []

    # --------------------------------------------------
    # Header (Name + Contact)
    # --------------------------------------------------
    name = header.get("name", "").strip() or "Your Name"
    email = header.get("email", "").strip()
    phone = header.get("phone", "").strip()
    linkedin = header.get("linkedin", "").strip()
    github = header.get("github", "").strip()
    location = header.get("location", "").strip()

    # Name big & bold
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(18)

    contact_parts = []
    if email:
        contact_parts.append(email)
    if phone:
        contact_parts.append(phone)
    if linkedin:
        contact_parts.append(f"LinkedIn: {linkedin}")
    if github:
        contact_parts.append(f"GitHub: {github}")
    if location:
        contact_parts.append(location)

    if contact_parts:
        p_contact = doc.add_paragraph(" | ".join(contact_parts))
        p_contact.style = doc.styles["Normal"]

    # Target role / headline line (optional)
    target_role = meta.get("targetRole", "").strip()
    if target_role:
        p_role = doc.add_paragraph(target_role)
        p_role.style = doc.styles["Normal"]

    # Small spacer
    doc.add_paragraph()

    # Helper for section heading
    def add_section_title(title: str):
        if not title:
            return
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(12)

    # --------------------------------------------------
    # Summary
    # --------------------------------------------------
    if summary.strip():
        add_section_title("Summary")
        doc.add_paragraph(summary.strip())
        doc.add_paragraph()  # spacer

    # --------------------------------------------------
    # Skills
    # --------------------------------------------------
    skills_clean = [s.strip() for s in skills if isinstance(s, str) and s.strip()]
    if skills_clean:
        add_section_title("Skills")
        for s in skills_clean:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(s)
        doc.add_paragraph()

    # --------------------------------------------------
    # Experience
    # --------------------------------------------------
    if experience:
        add_section_title("Experience")
        for exp in experience:
            title = (exp.get("title") or "").strip()
            company = (exp.get("company") or "").strip()
            loc = (exp.get("location") or "").strip()
            start = (exp.get("start") or "").strip()
            end = (exp.get("end") or "").strip()
            bullets = exp.get("bullets", []) or []

            header_parts = []
            if title:
                header_parts.append(title)
            if company:
                header_parts.append(company)
            header_line = " – ".join(header_parts) if header_parts else ""

            sub_parts = []
            if loc:
                sub_parts.append(loc)
            if start or end:
                sub_parts.append(f"{start} – {end}".strip(" –"))
            sub_line = " | ".join(sub_parts) if sub_parts else ""

            if header_line:
                doc.add_paragraph(header_line)
            if sub_line:
                p = doc.add_paragraph(sub_line)
                p.style = doc.styles["Normal"]

            for b in bullets:
                b = (b or "").strip()
                if not b:
                    continue
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(b)

            doc.add_paragraph()  # spacer between jobs

    # --------------------------------------------------
    # Education
    # --------------------------------------------------
    if education:
        add_section_title("Education")
        for edu in education:
            degree = (edu.get("degree") or "").strip()
            inst = (edu.get("institution") or "").strip()
            bullets = edu.get("bullets", []) or []

            line = " – ".join(p for p in [degree, inst] if p)
            if line:
                doc.add_paragraph(line)

            for b in bullets:
                b = (b or "").strip()
                if not b:
                    continue
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(b)

            doc.add_paragraph()

    # --------------------------------------------------
    # Projects
    # --------------------------------------------------
    if projects:
        add_section_title("Projects")
        for proj in projects:
            name_p = (proj.get("name") or "").strip()
            bullets = proj.get("bullets", []) or []

            if name_p:
                p = doc.add_paragraph()
                r = p.add_run(name_p)
                r.bold = True

            for b in bullets:
                b = (b or "").strip()
                if not b:
                    continue
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(b)

            doc.add_paragraph()

    # --------------------------------------------------
    # Extras
    # --------------------------------------------------
    extras_clean = [e.strip() for e in extras if isinstance(e, str) and e.strip()]
    if extras_clean:
        add_section_title("Additional")
        for e in extras_clean:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(e)

    # --------------------------------------------------
    # Save
    # --------------------------------------------------
    doc.save(str(out_path))
